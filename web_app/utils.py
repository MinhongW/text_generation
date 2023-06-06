import pandas as pd
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import streamlit as st
import time


def convert_text_to_table(text, separator=','):
    # Split the text by the separator
    data = text.split(separator)

    # Find the index of the first numeric value
    header_end_index = len(text.split(", ")[0].split(","))

    # Extract the header and reshape the data into a 2D list
    header, data = data[:header_end_index], data[header_end_index:]
    num_columns = len(header)
    data = [data[i:i+num_columns] for i in range(0, len(data), num_columns)]

    # Convert the data into a pandas DataFrame
    df = pd.DataFrame(data, columns=header)
    return df


def parse_input_string(user_input):
    if len(user_input) > 5:
        # Extract the caption
        # caption_start = user_input.index("caption:") + len("caption:")
        # caption_end = user_input.index(".table:")
        # caption = user_input[caption_start:caption_end].strip()
        #
        # # Extract the table data
        # table_start = user_input.index(".table:") + len(".table:")
        # table_text = user_input[table_start:].strip()
        #
        # df = convert_text_to_table(table_text)

        # Split the string into caption and table parts
        caption_start = user_input.find("caption:") + len("caption:")
        caption_text = user_input[caption_start:].split(".")[0].strip()

        table_start = user_input.find("table:") + len("table:")
        table_text = user_input[table_start:].strip()

        # Split the table text into rows
        rows = [row.strip().split(",") for row in table_text.split("\n")]

        # Create a DataFrame from the table data
        df = pd.DataFrame(rows[1:], columns=rows[0])

        model_input = user_input.replace("\n", ",")

        return caption_text, df, model_input

    else:
        return None, None


# def extract_table_as_df(table):
#     data = [[cell.text for cell in row.cells] for row in table.rows]
#     df = pd.DataFrame(data)
#     return df
#
#
# def read_docx_tables(file, tab_id=None):
#     document = Document(file)
#     if tab_id is None:
#         return [extract_table_as_df(table) for table in document.tables]
#     else:
#         return extract_table_as_df(document.tables[tab_id])


def extract_tables_from_docx(file):
    """Extracts tables from a Word document and returns them as pandas DataFrames."""
    document = Document(file)
    tables = []
    captions = []
    caption = None

    # Traverse all body elements
    for elem in document.element.body:
        # If the element is a table, convert it to a DataFrame and store the latest caption
        if elem.tag.endswith('tbl'):
            table = Table(elem, document)
            data = [[cell.text for cell in row.cells] for row in table.rows]
            df = pd.DataFrame(data)
            tables.append(df)
            captions.append(caption)
            caption = None  # Reset the caption for the next table
        # If the element is a paragraph and it's not empty, store it as a potential caption
        elif elem.tag.endswith('p'):
            para = Paragraph(elem, document)
            if para.text.strip():
                caption = para.text

    return captions, tables


def convert_docx_table_to_model_input(caption, table_df):
    table_vals = '\n'.join(table_df.apply(lambda row: ', '.join(row.astype(str)), axis=1))
    return "caption: " + caption + ". table:" + table_vals


# def display_text_word_by_word(text, delay=0.2, words_per_line=10):
#     placeholder = st.empty()
#     words = text.split()
#     for i in range(len(words)):
#         current_line = ' '.join(words[i-words_per_line+1:i+1]) if i+1 > words_per_line else ' '.join(words[:i+1])
#         previous_lines = '<br>'.join([' '.join(words[j:j+words_per_line]) for j in range(0, i, words_per_line)])
#         placeholder.markdown(f'{previous_lines}<br>{current_line}', unsafe_allow_html=True)
#         time.sleep(delay)


def display_text_word_by_word(input_text):
    """
    By Hassan
    API to print out text gradually, i.e., word by word
    """
    try:
        # Split the generated text into individual words
        words = input_text.split()
        # ToDo: Remove any exccessive white spaces

        # Not the most ideal solution, but it works ;)
        # The idea is that we are going to empty the text area, add a new word
        # and print out the new text. We can also use JavaScript but it's no necessary.
        empty_text = st.empty()
        output_text = ""
        for word in words:
            output_text += f"{word} "
            empty_text.markdown(output_text, unsafe_allow_html=True)
            time.sleep(0.1)  # Adjust the delay time as desired
    except Exception as e:
        raise (f'Ooops! {e}')

# def display_text_word_by_word(text, delay=0.2):
#     placeholder = st.empty()
#     words = text.split()
#     for i in range(len(words)):
#         placeholder.text(' '.join(words[:i+1]))
#         time.sleep(delay)

def generate_table_labels(total_tables):
    table_labels = ['Table {}'.format(i) for i in range(1, total_tables + 1)]
    return table_labels


def table_to_index(table_label):
    index = int(table_label.split(' ')[-1]) - 1
    return index
